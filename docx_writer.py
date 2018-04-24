from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    # check for tag existence, if none found, then create one
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existence, if none found, then create one
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


class DocxWriter:
    def __init__(self):
        self.document = Document('style.docx')

    def add_paragraph(self, text):
        p = self.document.add_paragraph(text)
        p.style = 'Normal'

    def add_heading(self, text, level):
        h = self.document.add_heading(text, level)
        h.style = 'Heading ' + str(level)

    def add_table(self, table, is_headers):
        new_table = self.document.add_table(rows=len(table), cols=len(table[0]))
        # new_table.style = 'Table'
        for i in range(0, len(table)):
            for j in range(0, len(table[0])):
                new_table.rows[i].cells[j].text = table[i][j]
        if is_headers:
            for cell in new_table.rows[0].cells:
                set_cell_border(cell, bottom={"sz": 1, "color": "#000000", "val": "single"})

    def add_bullet_element(self, string, level):
        p = self.document.add_paragraph()
        p.style = 'List Bullet'
        if level in range(2, 5):
            p.style = 'List Bullet ' + str(level)
        p.add_run(string)

    def save(self, path):
        self.document.save(path)
